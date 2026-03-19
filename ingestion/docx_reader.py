"""
DOCX Reader — Uses MarkItDown to convert Word documents to Markdown,
preserving heading hierarchy, bullet lists, numbered lists, and indentation.

MarkItDown (by Microsoft) converts .docx → clean Markdown:
  - Headings: # H1, ## H2, ### H3
  - Bullets:  - item / • item
  - Numbered: 1. item
  - Bold:     **text**

This module then parses the Markdown into structured sections matching
the original Word document hierarchy, ready for downstream chunking.
"""
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from src.core.logger import get_logger

logger = get_logger(__name__)

# ── Program level patterns (unchanged) ─────────────────────────────
LEVEL_PATTERNS = {
    "thac_si": [
        r"(?i)th[aạ]c\s*s[iĩỹ]",
        r"(?i)\bThS\b",
        r"(?i)master",
    ],
    "tien_si": [
        r"(?i)ti[eế]n\s*s[iĩỹ]",
        r"(?i)\bNCS\b",
        r"(?i)\bTS\b",
        r"(?i)doctor",
        r"(?i)ph\.?d",
    ],
}

# ── Program name extraction patterns ───────────────────────────────
PROGRAM_TITLE_PATTERN = re.compile(
    r"(?i)ch[uư][oơ]ng\s+tr[iì]nh\s+"
    r"(?:[đd][àa]o\s+t[aạ]o\s+)?"
    r"(?:th[aạ]c\s*s[iĩỹ]|ti[eế]n\s*s[iĩỹ])\s+"
    r"(?:ng[àa]nh\s+)?(.+)",
    re.UNICODE,
)
INTRO_TITLE_PATTERN = re.compile(
    r"(?i)gi[oớ]i\s+thi[eệ]u\s+ch[uư][oơ]ng\s+tr[iì]nh\s+"
    r"(?:[đd][àa]o\s+t[aạ]o\s+)?"
    r"(?:th[aạ]c\s*s[iĩỹ]|ti[eế]n\s*s[iĩỹ])\s+"
    r"(?:ng[àa]nh\s+)?(.+)",
    re.UNICODE,
)
INTRO_NGANH_PATTERN = re.compile(
    r"(?i)gi[oớ]i\s+thi[eệ]u\s+ch[uư][oơ]ng\s+tr[iì]nh\s+"
    r"(?:ng[àa]nh\s+)?(.+)",
    re.UNICODE,
)
FILENAME_PREFIX = re.compile(r"^(ThS|NCS|TS)\s+", re.IGNORECASE)
PROGRAM_NAME_CUTOFF = re.compile(
    r"\s+(?:tại|được|do|là|có|nhằm|theo|dựa)\s+",
    re.IGNORECASE | re.UNICODE,
)

# ── Markdown heading pattern ────────────────────────────────────────
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.+)$")

# ── Markdown inline cleanup (keep structure, strip markup syntax) ───
_MD_BOLD        = re.compile(r"\*\*(.+?)\*\*")
_MD_ITALIC      = re.compile(r"(?<![*])\*(?![*])(.+?)(?<![*])\*(?![*])")
_MD_INLINE_CODE = re.compile(r"`(.+?)`")
# [text](url) — standard Markdown link
_MD_LINK        = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
# <url> — MarkItDown autolink format for hyperlinks
_MD_AUTOLINK    = re.compile(r"<(https?://[^>\s]+)>")
# Bare URL — raw https:// on its own
_BARE_URL       = re.compile(r"https?://\S+")


@dataclass
class DocxSection:
    """A section extracted from a DOCX document."""
    heading_level: int   # 1–6 (Markdown heading depth)
    heading_text: str    # Clean heading without # prefix
    content: str         # Full text content, with bullets/numbers preserved
    paragraphs: list[str] = field(default_factory=list)


@dataclass
class DocxDocument:
    """Parsed DOCX document with structured sections."""
    filename: str
    program_name: str
    program_level: str   # "thac_si" | "tien_si"
    sections: list[DocxSection] = field(default_factory=list)
    raw_text: str = ""


class DocxReader:
    """
    Reads DOCX files using MarkItDown, preserving:
      - Heading hierarchy (H1–H6)
      - Bullet lists  (•, -, *)
      - Numbered lists (1. 2. 3. / a. b. c.)
      - Indentation levels
    """

    def read(self, file_path: Path) -> DocxDocument:
        """
        Read a DOCX or pre-converted TXT file → structured DocxDocument.

        - .txt files (processed/): read directly as Markdown text (faster, no MarkItDown)
        - .docx files (raw/): convert via MarkItDown + re-inject hyperlinks from XML
        """
        logger.info("docx_reader_start", file=str(file_path))

        suffix = file_path.suffix.lower()

        if suffix == ".txt":
            # Pre-converted Markdown text — read directly, no extra conversion needed
            markdown_text = file_path.read_text(encoding="utf-8")
        else:
            # .docx — convert with MarkItDown, then re-inject lost hyperlinks
            try:
                from markitdown import MarkItDown
            except ImportError:
                raise ImportError(
                    "markitdown is required. Install with: pip install markitdown"
                )
            md = MarkItDown()
            result = md.convert(str(file_path))
            markdown_text = result.text_content

            # Re-inject hyperlinks dropped by MarkItDown
            hyperlinks = self._extract_hyperlinks(file_path)
            if hyperlinks:
                markdown_text = self._inject_hyperlinks(markdown_text, hyperlinks)

        filename = file_path.name

        # Detect program level from filename or content
        program_level = self._detect_program_level(filename, markdown_text)

        # Parse markdown → sections
        sections = self._parse_markdown_sections(markdown_text)

        # Extract program name
        program_name = self._extract_program_name(sections, filename)

        doc = DocxDocument(
            filename=filename,
            program_name=program_name,
            program_level=program_level,
            sections=sections,
            raw_text=markdown_text,
        )

        logger.info(
            "docx_reader_done",
            file=filename,
            sections=len(sections),
            program=program_name,
            level=program_level,
        )
        return doc

    @staticmethod
    def _extract_hyperlinks(file_path: Path) -> dict[str, str]:
        """
        Extract all hyperlinks from a DOCX file via python-docx XML.

        Word stores hyperlinks in `document.xml.rels` as relationship entries.
        python-docx exposes them via `doc.part.rels`.

        Returns:
            {display_text: url} — e.g.
            {"Thông tin chi tiết chương trình đào tạo:": "https://daotaosdh.ufm.edu.vn/..."}
        """
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            return {}

        links: dict[str, str] = {}
        try:
            doc = Document(str(file_path))
            rels = doc.part.rels  # {rId: Relationship}

            # Walk all paragraphs
            for para in doc.paragraphs:
                for hyperlink_el in para._element.findall(".//{}".format(
                    qn("w:hyperlink")
                )):
                    r_id = hyperlink_el.get(qn("r:id"))
                    if not r_id or r_id not in rels:
                        continue

                    rel = rels[r_id]
                    # Only external URLs (HTTP/HTTPS)
                    url = rel.target_ref
                    if not url or not url.startswith("http"):
                        continue

                    # Build display text from all <w:t> children
                    texts = [
                        t.text for t in hyperlink_el.findall(
                            ".//{}".format(qn("w:t"))
                        ) if t.text
                    ]
                    display = "".join(texts).strip()
                    if display:
                        links[display] = url

        except Exception as e:
            logger.warning("hyperlink_extract_failed", error=str(e))

        return links

    @staticmethod
    def _inject_hyperlinks(markdown: str, hyperlinks: dict[str, str]) -> str:
        """
        Re-inject hyperlink URLs into MarkItDown markdown output.

        For each (display_text, url) pair:
        - If `display_text` appears in markdown WITHOUT a URL right after it,
          append ` (url)` immediately after the display text.
        - Handles display text that appears as a standalone line or mid-paragraph.
        - Skips injection if URL is already present in the line.
        """
        lines = markdown.splitlines()
        result_lines = []

        for line in lines:
            for display, url in hyperlinks.items():
                # Only inject if display text is in this line but URL is not yet present
                if display in line and url not in line:
                    # Replace first occurrence: add URL right after the display text
                    line = line.replace(display, f"{display} ({url})", 1)
            result_lines.append(line)

        return "\n".join(result_lines)

    # ────────────────────────────────────────────────────────────────
    # Markdown → Section Parser
    # ────────────────────────────────────────────────────────────────

    def _parse_markdown_sections(self, markdown: str) -> list[DocxSection]:
        """
        Parse Markdown text into DocxSections.

        Heading detection — TWO strategies (hybrid):
          1. Standard Markdown: lines matching `^#{1,6} heading`
          2. ALL-CAPS heuristic: lines that are ALL-CAPS (≥80% uppercase letters)
             OR wrapped in **BOLD-CAPS** — common in Vietnamese Word documents
             where section headings are bold+uppercase but not styled as Heading.

        Content lines (bullets, numbered items, plain text) go into current section.
        """
        lines = markdown.splitlines()

        sections: list[DocxSection] = []
        current_heading_text = ""
        current_heading_level = 0
        current_lines: list[str] = []

        def _flush():
            if not current_heading_text and not current_lines:
                return
            paras = self._lines_to_paragraphs(current_lines)
            content = "\n".join(paras)
            sections.append(DocxSection(
                heading_level=current_heading_level,
                heading_text=self._clean_inline(current_heading_text),
                content=content.strip(),
                paragraphs=paras,
            ))

        for line in lines:
            stripped = line.strip()

            # Strategy 1: Standard Markdown heading
            m = _MD_HEADING.match(line)
            if m:
                _flush()
                current_heading_level = len(m.group(1))
                current_heading_text = m.group(2).strip()
                current_lines = []
                continue

            # Strategy 2: ALL-CAPS or **BOLD-CAPS** implicit heading
            if self._is_all_caps_heading(stripped):
                _flush()
                current_heading_level = 1
                # Remove bold markers if present
                current_heading_text = re.sub(r'\*+', '', stripped).strip()
                current_lines = []
                continue

            # Regular content line
            current_lines.append(line)

        _flush()
        return sections

    @staticmethod
    def _is_all_caps_heading(text: str) -> bool:
        """
        Detect ALL-CAPS or **BOLD-CAPS** lines as implicit headings.

        Rules:
        - Length 4–150 chars
        - ≥80% of alphabetic chars are uppercase (Vietnamese caps included)
        - Does NOT start with a bullet symbol (not a list item)
        - Does NOT look like a URL
        - Has at least 2 words (not just a single ALL-CAPS word in body text)
        """
        if not text or len(text) < 4 or len(text) > 150:
            return False

        # Strip bold markers for analysis
        clean = re.sub(r'\*+', '', text).strip()
        if not clean:
            return False

        # Not a list item
        if clean[0] in '•–▸·-*+' or re.match(r'^\d+\.', clean):
            return False

        # Not a URL
        if clean.startswith('http'):
            return False

        # Must have at least 2 words (headings aren't single-word blurbs)
        words = clean.split()
        if len(words) < 2:
            return False

        # Check uppercase ratio on alphabetic chars
        alpha = [c for c in clean if c.isalpha()]
        if len(alpha) < 4:
            return False
        upper_ratio = sum(1 for c in alpha if c.isupper()) / len(alpha)
        return upper_ratio >= 0.80

    def _lines_to_paragraphs(self, lines: list[str]) -> list[str]:
        """
        Convert raw content lines into clean paragraphs.

        Rules:
        - Bullet lines (-, *, •, +) → keep the bullet symbol, clean inline markdown
        - Numbered lines (1., 2., a., i.) → keep the prefix, clean inline markdown
        - Blank lines → paragraph separator (collapsed)
        - Plain text → clean inline markdown
        - Consecutive blank lines → single blank
        """
        paragraphs = []
        current_para_lines: list[str] = []
        prev_blank = False

        for raw_line in lines:
            line = raw_line.rstrip()

            # Blank line → separator
            if not line.strip():
                if current_para_lines:
                    paragraphs.append(self._clean_inline(" ".join(current_para_lines)))
                    current_para_lines = []
                if not prev_blank and paragraphs:
                    paragraphs.append("")   # preserve blank separator
                prev_blank = True
                continue

            prev_blank = False

            # Bullet line: starts with */+/-/• possibly indented
            # MarkItDown uses * for bullets, + for sub-bullets
            bullet_match = re.match(
                r"^(\s*)([-*+•]|\d+\.|[a-zA-Z]+\.)\s+(.*)", line
            )
            if bullet_match:
                # Flush accumulated paragraph text first
                if current_para_lines:
                    paragraphs.append(self._clean_inline(" ".join(current_para_lines)))
                    current_para_lines = []

                indent        = bullet_match.group(1)
                bullet_marker = bullet_match.group(2)
                item_text     = self._clean_inline(bullet_match.group(3))

                # Normalize bullet symbols:
                #   indent 0       → • (main bullet or 1. for numbered)
                #   indent 1–2     → – (sub-item)
                #   indent 3–4     → ▸ (sub-sub-item)
                #   indent 5+      → · (deep nested)
                indent_len = len(indent)
                if re.match(r"\d+\.|[a-zA-Z]+\.", bullet_marker):
                    # Numbered list — keep as "1. " prefix
                    prefix = f"{indent}{bullet_marker} "
                else:
                    # Bullet list — normalize symbol by depth
                    symbols = ["• ", "– ", "▸ ", "· "]
                    depth = min(indent_len // 2, len(symbols) - 1)
                    prefix = f"{indent}{symbols[depth]}"

                paragraphs.append(f"{prefix}{item_text}")
            else:
                stripped_line = line.strip()
                # Bare URL on its own → flush accumulated text first, then output URL as standalone para
                if _BARE_URL.search(stripped_line) and len(stripped_line) < 300:
                    if current_para_lines:
                        paragraphs.append(self._clean_inline(" ".join(current_para_lines)))
                        current_para_lines = []
                    paragraphs.append(self._clean_inline(stripped_line))
                else:
                    # Regular text line — accumulate to join with neighbors
                    current_para_lines.append(stripped_line)

        # Flush remaining
        if current_para_lines:
            paragraphs.append(self._clean_inline(" ".join(current_para_lines)))

        # Remove leading/trailing blanks
        while paragraphs and paragraphs[0] == "":
            paragraphs.pop(0)
        while paragraphs and paragraphs[-1] == "":
            paragraphs.pop()

        return paragraphs

    @staticmethod
    def _clean_inline(text: str) -> str:
        """
        Strip Markdown inline markup but KEEP URLs and list structure.

        Handles:
          **bold** → bold
          *italic*  → italic  (but not bullet symbols)
          `code`    → code
          [text](url) → text (url)
          <https://...> → https://...  (MarkItDown autolink format)
        """
        # <url> autolink (MarkItDown format for hyperlinks)
        text = _MD_AUTOLINK.sub(r"\1", text)
        # **bold** → bold
        text = _MD_BOLD.sub(r"\1", text)
        # *italic* → italic  (careful not to eat bullet `* item`)
        text = _MD_ITALIC.sub(r"\1", text)
        # `code` → code
        text = _MD_INLINE_CODE.sub(r"\1", text)
        # [text](url) → text (url)
        text = _MD_LINK.sub(r"\1 (\2)", text)
        # collapse extra whitespace
        text = re.sub(r"[ \t]+", " ", text).strip()
        return text

    # ────────────────────────────────────────────────────────────────
    # Program name & level detection (same logic as before)
    # ────────────────────────────────────────────────────────────────

    def _detect_program_level(self, filename: str, content: str) -> str:
        """Detect program level from filename or markdown content."""
        for level, patterns in LEVEL_PATTERNS.items():
            for pat in patterns:
                if re.search(pat, filename):
                    return level
        # Check first 500 chars of content
        snippet = content[:500]
        for level, patterns in LEVEL_PATTERNS.items():
            for pat in patterns:
                if re.search(pat, snippet):
                    return level
        return "unknown"

    def _extract_program_name(
        self, sections: list[DocxSection], filename: str
    ) -> str:
        """
        Extract clean program name.
        Priority:
          1. PROGRAM_TITLE_PATTERN  from any heading
          2. INTRO_TITLE_PATTERN    from heading
          3. INTRO_NGANH_PATTERN    from heading
          4. Filename fallback
        """
        for section in sections:
            h = section.heading_text.strip()
            if not h:
                continue
            m = PROGRAM_TITLE_PATTERN.match(h)
            if m:
                return self._clean_program_name(m.group(1))
            m = INTRO_TITLE_PATTERN.match(h)
            if m:
                return self._clean_program_name(m.group(1))

        for section in sections:
            h = section.heading_text.strip()
            m = INTRO_NGANH_PATTERN.match(h)
            if m:
                return self._clean_program_name(m.group(1))

        # Fallback: derive from filename
        name = Path(filename).stem
        name = FILENAME_PREFIX.sub("", name).strip()
        return name

    @staticmethod
    def _clean_program_name(raw: str) -> str:
        """Trim trailing clauses like 'tại Trường...' from program name."""
        name = raw.strip()
        m = PROGRAM_NAME_CUTOFF.search(name)
        if m:
            name = name[:m.start()].strip()
        return name.rstrip(".,;:–-").strip()
